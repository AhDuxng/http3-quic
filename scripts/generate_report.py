"""
Script tạo báo cáo MetricCalculationReport.docx
Báo cáo chi tiết cách tính từng thông số đo lường cho paper Q1.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Tô nền cell trong bảng."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_metric_section(doc, number, title, how, how_code, dash_relation, meaning):
    """Thêm một mục thông số vào tài liệu."""
    doc.add_heading(f'{number}. {title}', level=2)

    # Lấy kiểu gì?
    p = doc.add_paragraph()
    run = p.add_run('Cách lấy dữ liệu: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)
    p.add_run(how)

    if how_code:
        code_p = doc.add_paragraph()
        code_p.style = doc.styles['No Spacing']
        code_p.paragraph_format.left_indent = Cm(1)
        code_p.paragraph_format.space_before = Pt(4)
        code_p.paragraph_format.space_after = Pt(4)
        for line in how_code.split('\n'):
            run = code_p.add_run(line + '\n')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)

    # Liên quan gì đến DASH?
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Liên quan đến DASH: ')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0, 100, 0)
    p2.add_run(dash_relation)

    # Ý nghĩa
    p3 = doc.add_paragraph()
    run3 = p3.add_run('Ý nghĩa: ')
    run3.bold = True
    run3.font.color.rgb = RGBColor(139, 0, 0)
    p3.add_run(meaning)

    doc.add_paragraph()  # spacing


def main():
    doc = Document()

    # ===== Page setup =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # ===== Default font =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 51, 102)

    # ===== TITLE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run('BÁO CÁO CHI TIẾT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run('CÁCH TÍNH TỪNG THÔNG SỐ ĐO LƯỜNG\nCHO NGHIÊN CỨU SO SÁNH HTTP/2 VÀ HTTP/3 TRONG DASH STREAMING')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(20)
    run = info.add_run('Hệ thống: ADTUBE Stream Analyzer\nCông nghệ: dash.js + PerformanceResourceTiming API + Network Information API')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ===== MỤC LỤC =====
    doc.add_heading('Mục lục', level=1)
    toc_items = [
        'Tổng quan hệ thống đo lường',
        'Chi tiết 16 thông số',
        '  2.1. Throughput (Throughput_kbps)',
        '  2.2. Time To First Byte — TTFB (TTFB_ms)',
        '  2.3. Segment Download Time — SDT (SDT_ms)',
        '  2.4. SDT Jitter (Jitter_ms)',
        '  2.5. Buffer Level (Buffer_s)',
        '  2.6. Stall Count (StallCount)',
        '  2.7. Stall Duration (StallDuration_ms)',
        '  2.8. Rebuffering Ratio (RebufferingRatio)',
        '  2.9. Video Bitrate (Bitrate_kbps)',
        '  2.10. Quality Switch Count (QualitySwitchCount)',
        '  2.11. Dropped Frames (DroppedFrames)',
        '  2.12. Frame Rate — FPS (FPS)',
        '  2.13. Download Speed (DownloadSpeed_kbps)',
        '  2.14. Resolution (Resolution)',
        '  2.15. Protocol (Protocol)',
        '  2.16. Network Type (NetworkType)',
        'Bảng tổng hợp',
        'Tài liệu tham khảo',
    ]
    for i, item in enumerate(toc_items):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if item.startswith('  '):
            p.paragraph_format.left_indent = Cm(1)
            p.add_run(item.strip()).font.size = Pt(11)
        else:
            num = i + 1 if i < 2 else i - 14 if i >= 18 else None
            label = f'{["1", "2", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "3", "4"][i]}. {item}' if not item.startswith('  ') else item
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(12)

    doc.add_page_break()

    # ===== PHẦN 1: TỔNG QUAN =====
    doc.add_heading('1. Tổng quan hệ thống đo lường', level=1)

    doc.add_paragraph(
        'Hệ thống ADTUBE Stream Analyzer thu thập 16 thông số đo lường từ 3 nguồn dữ liệu chính:'
    )

    # Bảng nguồn dữ liệu
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Nguồn dữ liệu', 'API / Sự kiện', 'Thông số thu được']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    sources = [
        ['Browser APIs', 'PerformanceResourceTiming\nVideoPlaybackQuality\nNetwork Information API', 'Protocol, TTFB, FPS,\nDroppedFrames, NetworkType'],
        ['dash.js Events', 'FRAGMENT_LOADING_COMPLETED\nBUFFER_EMPTY / BUFFER_LOADED\nQUALITY_CHANGE_RENDERED', 'SDT, Bytes, StallCount,\nStallDuration, QualitySwitchCount'],
        ['dash.js API', 'getBufferLength()\ngetAverageThroughput()\ngetRepresentationsByType()', 'Buffer, Throughput (fallback),\nBitrate, Resolution'],
    ]
    for r, row_data in enumerate(sources):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Luồng xử lý: Khi dash.js tải xong một segment (FRAGMENT_LOADING_COMPLETED), hệ thống tính toán SDT, TTFB, '
        'Jitter, Download Speed. Mỗi giây (polling 1s), hệ thống cập nhật Buffer, FPS, Throughput, Protocol, '
        'RebufferingRatio. Tất cả được ghi vào snapshot kèm theo mỗi dòng log, và xuất ra CSV 23 cột.'
    )

    doc.add_paragraph(
        'Thứ tự 23 cột CSV:'
    ).runs[0].bold = True

    csv_p = doc.add_paragraph()
    csv_p.paragraph_format.left_indent = Cm(1)
    run = csv_p.add_run(
        'Timestamp, Level, Message, Protocol, NetworkType, Bitrate_kbps, Resolution, '
        'Throughput_kbps, Buffer_s, FPS, TTFB_ms, SDT_ms, Jitter_ms, DownloadSpeed_kbps, '
        'StallCount, StallDuration_ms, RebufferingRatio, DroppedFrames, QualitySwitchCount, '
        'CurrentTime_s, Duration_s, IsAutoQuality, ActiveScenario'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===== PHẦN 2: CHI TIẾT 16 THÔNG SỐ =====
    doc.add_heading('2. Chi tiết 16 thông số đo lường', level=1)

    doc.add_paragraph(
        'Mỗi thông số được trình bày theo 3 câu hỏi: Lấy dữ liệu kiểu gì? Liên quan gì đến DASH? Ý nghĩa là gì?'
    )

    # ----- 2.1 Throughput -----
    add_metric_section(doc,
        '2.1', 'Throughput (Throughput_kbps)',
        'Mỗi khi tải xong một segment (sự kiện FRAGMENT_LOADING_COMPLETED), hệ thống tính tốc độ tải: '
        '(bytesLoaded × 8) / SDT và lưu vào mảng mẫu sliding window 10 giây. '
        'Khi polling (mỗi 1 giây), lấy trung bình cộng các mẫu trong 1 giây gần nhất. '
        'Nếu không có mẫu, fallback sang player.getAverageThroughput("video").',
        'File: useStreamMetrics.ts — hàm pollStats()\nThroughput = average(segment_speeds trong 1s gần nhất)\nFallback: player.getAverageThroughput("video")',
        'DASH chia video thành nhiều segment nhỏ (thường 2-4 giây). Mỗi segment tải xuống độc lập qua HTTP. '
        'Throughput ở đây đo sức chứa thực tế của mạng khi DASH player tải các segment — '
        'khác với throughput lý thuyết của đường truyền. '
        'Thuật toán ABR (Adaptive Bitrate) dựa vào Throughput để quyết định chọn representation nào cho segment tiếp theo.',
        'Cao hơn = tốt hơn. Throughput cao nghĩa là mạng có khả năng truyền dữ liệu nhanh → '
        'player chọn được representation bitrate cao hơn → chất lượng video tốt hơn. '
        'Đây là chỉ số then chốt khi so sánh H2 vs H3: QUIC (H3) được kỳ vọng có throughput ổn định hơn '
        'nhờ multiplexing không bị head-of-line blocking và 0-RTT handshake.'
    )

    # ----- 2.2 TTFB -----
    add_metric_section(doc,
        '2.2', 'Time To First Byte — TTFB (TTFB_ms)',
        'Sử dụng PerformanceResourceTiming API của browser. Tìm entry khớp URL segment, '
        'tính TTFB = entry.responseStart - entry.requestStart. '
        'Yêu cầu server phải có header Timing-Allow-Origin để browser cung cấp dữ liệu timing chính xác. '
        'Fallback: dùng firstByteDate - startDate từ đối tượng request của dash.js.',
        'File: performanceApi.ts — hàm getTTFBFromPerformanceAPI()\nTTFB = entry.responseStart - entry.requestStart\nFallback: req.firstByteDate - req.startDate',
        'Mỗi segment DASH là một HTTP request riêng. TTFB đo thời gian từ khi browser gửi request '
        'đến khi nhận byte đầu tiên của response. Trong bối cảnh so sánh H2 vs H3:\n'
        '• HTTP/2: Phải chờ TLS + TCP handshake (nếu kết nối mới), hoặc multiplex trên kết nối TCP sẵn có.\n'
        '• HTTP/3 (QUIC): Có thể dùng 0-RTT handshake → TTFB thấp hơn đáng kể ở connection đầu tiên.\n'
        'Lưu ý quan trọng: TTFB KHÔNG phải RTT (Round Trip Time). RTT không đo chính xác được từ browser. '
        'TTFB bao gồm: DNS + TCP/TLS handshake + server processing + 1 × network propagation.',
        'Thấp hơn = tốt hơn. TTFB thấp nghĩa là server phản hồi nhanh. '
        'Đây là chỉ số khác biệt lớn nhất giữa H2 và H3 — QUIC giảm được 1 RTT trong quá trình handshake. '
        'TTFB là chỉ số bắt buộc trong mọi paper so sánh giao thức truyền tải.'
    )

    # ----- 2.3 SDT -----
    add_metric_section(doc,
        '2.3', 'Segment Download Time — SDT (SDT_ms)',
        'Tính tổng thời gian tải một segment DASH. Hệ thống thử 3 nguồn theo thứ tự ưu tiên:\n'
        '1) req.endDate - req.startDate (dash.js v5+ FragmentRequest)\n'
        '2) Tổng req.trace[].d (mảng trace của dash.js)\n'
        '3) entry.responseEnd - entry.requestStart (Performance Resource Timing API)',
        'File: useStreamMetrics.ts — hàm processSegment()\nSDT = endDate - startDate (ưu tiên)\nFallback 1: sum(trace[].d)\nFallback 2: responseEnd - requestStart (Performance API)',
        'SDT là tổng thời gian tải một segment DASH. Trong DASH, player quyết định chọn bitrate dựa trên SDT '
        'của các segment trước đó:\n'
        '• SDT ngắn → mạng nhanh → thuật toán ABR chọn bitrate cao hơn\n'
        '• SDT dài → mạng chậm → ABR giảm bitrate để tránh stall\n'
        'SDT liên quan trực tiếp đến Throughput (Throughput = bytes/SDT) và Jitter.',
        'Thấp hơn = tốt hơn. SDT ngắn nghĩa là segment tải nhanh → ít bị ngắt buffer. '
        'SDT ổn định (jitter thấp) giúp ABR chọn bitrate chính xác hơn. '
        'Trong paper, so sánh mean SDT và phân phối SDT giữa H2 và H3 cho thấy giao thức nào tải segment nhanh hơn.'
    )

    # ----- 2.4 Jitter -----
    add_metric_section(doc,
        '2.4', 'SDT Jitter (Jitter_ms)',
        'Tính sau mỗi segment tải xong. Giữ SDT trước đó trong biến prevSDTRef. '
        'Jitter = |SDT hiện tại - SDT trước đó|. Nếu là segment đầu tiên (chưa có SDT trước), jitter = 0.',
        'File: useStreamMetrics.ts — hàm processSegment()\nJitter = |SDT_hiện_tại - SDT_trước_đó|',
        'Jitter cao nghĩa là thời gian tải segment dao động lớn → thuật toán ABR khó dự đoán throughput → '
        'chọn sai bitrate → ảnh hưởng trải nghiệm (nhảy quality hoặc stall). '
        'H3 (QUIC) được kỳ vọng có jitter thấp hơn H2 trong môi trường mất gói tin (packet loss) '
        'nhờ cơ chế khôi phục gói nhanh hơn và không bị head-of-line blocking ở tầng transport.',
        'Thấp hơn = tốt hơn. Jitter thấp → mạng ổn định → ABR chọn bitrate chính xác hơn → '
        'ít quality switch và ít stall. '
        'Trong paper, biểu đồ CDF (Cumulative Distribution Function) của Jitter giữa H2 và H3 '
        'là hình ảnh thuyết phục cho thấy sự khác biệt về tính ổn định mạng.'
    )

    # ----- 2.5 Buffer -----
    add_metric_section(doc,
        '2.5', 'Buffer Level (Buffer_s)',
        'Gọi trực tiếp API của dash.js player mỗi 1 giây (polling interval).',
        'File: useStreamMetrics.ts — hàm pollStats()\nbufferSeconds = player.getBufferLength("video")',
        'Buffer là lượng dữ liệu video đã tải xuống nhưng chưa phát. Trong DASH:\n'
        '• Player tải segment → dữ liệu vào buffer\n'
        '• Video phát → dữ liệu ra khỏi buffer\n'
        '• Nếu buffer = 0 → video dừng hình (stall)\n'
        'Buffer level là yếu tố quyết định để ABR chọn tăng hay giảm bitrate.',
        'Cao hơn = tốt hơn (ít rủi ro stall). Buffer thấp (<2s) là dấu hiệu nguy hiểm. '
        'Trong paper, so sánh average buffer level giữa H2 và H3 cho thấy giao thức nào duy trì buffer tốt hơn, '
        'đặc biệt trong điều kiện mạng xấu.'
    )

    # ----- 2.6 StallCount -----
    add_metric_section(doc,
        '2.6', 'Stall Count (StallCount)',
        'Mỗi khi dash.js phát sự kiện BUFFER_EMPTY (mediaType === "video"), '
        'hệ thống tăng bộ đếm stallCountRef lên 1. '
        'Chỉ đếm sự kiện BUFFER_EMPTY từ dash.js — KHÔNG dùng sự kiện "waiting" của HTML5 '
        '(vì "waiting" có thể fire vì nhiều lý do khác như seek, initial load).',
        'File: useStallTracker.ts — hàm onBufferEmpty()\nMỗi BUFFER_EMPTY event: stallCount += 1',
        'Stall = sự kiện buffer bị cạn (hết dữ liệu để phát). Xảy ra khi tốc độ tải segment < tốc độ phát video. '
        'Trong DASH, đây là hậu quả của:\n'
        '• Mạng quá chậm so với bitrate đang phát\n'
        '• ABR chọn bitrate quá cao\n'
        '• Môi trường mạng đột ngột xấu đi',
        'Thấp hơn = tốt hơn. Mỗi stall là một lần video bị dừng hình → trải nghiệm người dùng rất xấu. '
        'Đây là chỉ số quan trọng nhất trong QoE research. Paper Q1 bắt buộc phải có. '
        'BUFFER_EMPTY (dash.js) chính xác hơn HTML5 "waiting" event vì chỉ phản ánh buffer depletion thực sự.'
    )

    # ----- 2.7 StallDuration -----
    add_metric_section(doc,
        '2.7', 'Stall Duration (StallDuration_ms)',
        'Khi nhận BUFFER_EMPTY: ghi lại thời điểm bắt đầu = Date.now(). '
        'Khi nhận BUFFER_LOADED: tính duration = Date.now() - thời điểm bắt đầu, '
        'cộng dồn vào stallAccumulatedMs. Đây là tổng tích lũy trong toàn phiên.',
        'File: useStallTracker.ts\nBUFFER_EMPTY → ghi startTime\nBUFFER_LOADED → duration = now - startTime\nstallAccumulatedMs += duration',
        'Tổng thời gian video bị dừng hình do buffer cạn. Mỗi lần BUFFER_EMPTY → BUFFER_LOADED là một chu kỳ stall.',
        'Thấp hơn = tốt hơn. Thời gian stall dài nghĩa là người xem phải chờ lâu → suy giảm QoE nghiêm trọng. '
        'Kết hợp với StallCount để tính trung bình thời gian mỗi lần stall: StallDuration / StallCount. '
        'H3 với cơ chế khôi phục nhanh hơn được kỳ vọng có stall duration ngắn hơn H2.'
    )

    # ----- 2.8 RebufferingRatio -----
    add_metric_section(doc,
        '2.8', 'Rebuffering Ratio (RebufferingRatio)',
        'Tính mỗi giây trong hàm polling.',
        'File: useStreamMetrics.ts — hàm pollStats()\nRebufferingRatio = stallAccumulatedMs / (currentTime × 1000)\n\n• stallAccumulatedMs: tổng thời gian stall tích lũy (từ useStallTracker)\n• currentTime: vị trí phát hiện tại (giây) × 1000 → ms',
        'Tỷ lệ thời gian người xem phải chờ do buffer cạn so với tổng thời gian xem video. '
        'Đây là chỉ số QoE tiêu chuẩn được định nghĩa bởi Seufert et al. (IEEE Communications Surveys & Tutorials, 2015). '
        'Giá trị nằm trong khoảng [0, 1].',
        'Thấp hơn = tốt hơn. Mức chấp nhận cho streaming: < 1% (0.01). '
        'Ví dụ: RebufferingRatio = 0.02 nghĩa là 2% thời gian xem bị dừng hình. '
        'Đây là chỉ số bắt buộc trong mọi paper QoE về adaptive streaming. '
        'So sánh RebufferingRatio giữa H2 và H3 là trọng tâm của bài nghiên cứu.'
    )

    # ----- 2.9 Bitrate -----
    add_metric_section(doc,
        '2.9', 'Video Bitrate (Bitrate_kbps)',
        'Lấy trực tiếp từ dash.js API.',
        'File: useStreamMetrics.ts — hàm getRepBitrateKbps()\nƯu tiên: rep.bitrateInKbit (dash.js v5+)\nFallback: Math.round(rep.bandwidth / 1000)\n\nrep = player.getCurrentRepresentationForType("video")',
        'Trong DASH, video được mã hóa ở nhiều mức bitrate (representations). '
        'Thuật toán ABR chọn representation phù hợp dựa trên throughput và buffer hiện tại. '
        'Bitrate phản ánh chất lượng hình ảnh trực tiếp — bitrate cao = nhiều chi tiết, ít artifact nén.',
        'Cao hơn = tốt hơn (chất lượng video tốt hơn). '
        'So sánh H2 vs H3: giao thức nào cho phép duy trì average bitrate cao hơn trong cùng điều kiện mạng. '
        'Biểu đồ time-series Bitrate theo thời gian là hình ảnh bắt buộc trong paper.'
    )

    # ----- 2.10 QualitySwitchCount -----
    add_metric_section(doc,
        '2.10', 'Quality Switch Count (QualitySwitchCount)',
        'Mỗi khi dash.js phát sự kiện QUALITY_CHANGE_RENDERED (mediaType === "video"), '
        'hệ thống tăng bộ đếm lên 1.',
        'File: useStreamMetrics.ts — hàm incrementQualitySwitch()\nMỗi QUALITY_CHANGE_RENDERED event: count += 1',
        'ABR algorithm của DASH chuyển đổi giữa các representation khi throughput thay đổi. '
        'Mỗi lần chuyển = 1 quality switch. Chuyển nhiều gây hiện tượng "nhảy hình" (flickering).',
        'Thấp hơn = tốt hơn. Chuyển quality nhiều → người xem thấy video liên tục thay đổi độ sắc nét → khó chịu. '
        'H3 với throughput ổn định hơn sẽ có ít quality switch hơn H2 trong môi trường mạng xấu. '
        'Trong paper, nên trình bày dưới dạng boxplot để thấy phân phối.'
    )

    # ----- 2.11 DroppedFrames -----
    add_metric_section(doc,
        '2.11', 'Dropped Frames (DroppedFrames)',
        'Lấy trực tiếp từ Browser API.',
        'File: useStreamMetrics.ts — hàm pollStats()\ndroppedFrames = video.getVideoPlaybackQuality()?.droppedVideoFrames ?? 0',
        'Frame bị rơi khi CPU/GPU không decode kịp. Trong DASH, nếu ABR chọn bitrate quá cao '
        'cho thiết bị, sẽ có nhiều dropped frames. Liên quan đến khả năng decode của thiết bị nhiều hơn giao thức.',
        'Thấp hơn = tốt hơn. Frame bị rơi gây giật hình. '
        'Chỉ số này ít liên quan trực tiếp đến so sánh H2 vs H3 (cùng video codec), '
        'nhưng nên ghi nhận để loại trừ ảnh hưởng của thiết bị trong kết quả thí nghiệm.'
    )

    # ----- 2.12 FPS -----
    add_metric_section(doc,
        '2.12', 'Frame Rate — FPS (FPS)',
        'Tính từ VideoPlaybackQuality API mỗi giây.',
        'File: useStreamMetrics.ts — hàm pollStats()\nFPS = (totalFrames_t2 - totalFrames_t1) / (currentTime_t2 - currentTime_t1)\n\n• totalFrames từ VideoPlaybackQuality.totalVideoFrames\n• currentTime từ HTMLVideoElement.currentTime\n• Tính mỗi 1 giây (polling interval)',
        'FPS cho thấy tốc độ render thực tế. Khi có stall hoặc dropped frame, FPS giảm — '
        'cho thấy video không mượt. Giá trị FPS bình thường cho video 30fps nên gần 30.',
        'Cao hơn = tốt hơn. FPS giảm nhiều so với framerate gốc → có vấn đề decode hoặc buffer. '
        'Lưu ý quan trọng: FPS được lưu dạng số (number), KHÔNG phải chuỗi (string) — '
        'để script Python có thể tính mean/median/std chính xác.'
    )

    # ----- 2.13 DownloadSpeed -----
    add_metric_section(doc,
        '2.13', 'Download Speed (DownloadSpeed_kbps)',
        'Tính ngay sau mỗi segment tải xong.',
        'File: useStreamMetrics.ts — hàm processSegment()\nDownloadSpeed = (bytesLoaded × 8) / SDT_ms\n\n• bytesLoaded: kích thước segment (bytes)\n• SDT_ms: thời gian tải segment (ms)\n• Kết quả: kbps (kilobits per second)',
        'Tốc độ tải segment cụ thể (khác Throughput là trung bình nhiều segment). '
        'Dùng để đánh giá hiệu năng mạng tại thời điểm cụ thể.',
        'Cao hơn = tốt hơn. Phản ánh tốc độ mạng thực tại thời điểm tải 1 segment cụ thể. '
        'Trong paper, biểu đồ scatter plot Download Speed theo thời gian giúp thấy sự dao động.'
    )

    # ----- 2.14 Resolution -----
    add_metric_section(doc,
        '2.14', 'Resolution (Resolution)',
        'Lấy trực tiếp từ representation hiện tại của dash.js.',
        'File: useStreamMetrics.ts — hàm getResolutionLabel()\nresolutionLabel = `${rep.width}x${rep.height}`\nrep = player.getCurrentRepresentationForType("video")',
        'Mỗi representation DASH có kích thước khung hình khác nhau (ví dụ: 640×360, 1280×720, 1920×1080). '
        'ABR chọn representation → resolution tương ứng.',
        'Metadata, không phải metric so sánh trực tiếp H2 vs H3 '
        '(vì cùng giao thức khác nhau nhưng cùng video sẽ có cùng resolution choices). '
        'Hữu ích để biết player đang phát ở mức nào tại mỗi thời điểm.'
    )

    # ----- 2.15 Protocol -----
    add_metric_section(doc,
        '2.15', 'Protocol (Protocol)',
        'Sử dụng PerformanceResourceTiming API.',
        'File: performanceApi.ts — hàm detectProtocol()\n1) Lấy entries có URL chứa "/media"\n2) Đọc entry.nextHopProtocol:\n   • "h3" hoặc "h3-29" → "HTTP/3 (QUIC)"\n   • "h2" → "HTTP/2"\n   • "http/1.1" → "HTTP/1.1"\n3) Fallback: tìm trên tất cả entries\n4) Mặc định: "Detecting..."',
        'Đây là biến độc lập chính của bài nghiên cứu. Xác định segment DASH được tải qua giao thức nào. '
        'Mỗi dòng CSV ghi lại protocol tại thời điểm đó → có thể phân tích hiện tượng "fallback" (H3 → H2) khi mạng xấu.',
        'Phân biệt H2 và H3 trong dữ liệu — cơ bản để chia nhóm và so sánh. '
        'Đây là biến phân loại chính (independent variable) của thí nghiệm.'
    )

    # ----- 2.16 NetworkType -----
    add_metric_section(doc,
        '2.16', 'Network Type (NetworkType)',
        'Sử dụng Network Information API — thuộc tính connection.type (loại kết nối vật lý thực tế).',
        'File: performanceApi.ts — hàm getNetworkType()\nnetworkType = navigator.connection.type\nGiá trị: "wifi", "cellular", "ethernet", "none", "unknown"',
        'Xác định môi trường mạng thực tế của phiên đo lường. Trong paper, cần ghi rõ test qua WiFi hay Cellular '
        '→ đảm bảo reproducibility. '
        'LƯU Ý QUAN TRỌNG: KHÔNG dùng navigator.connection.effectiveType — '
        'giá trị này luôn trả "4g" cho mọi kết nối WiFi tốt vì nó ước tính chất lượng mạng, '
        'KHÔNG phải loại kết nối vật lý. Đây là lỗi phổ biến trong nhiều hệ thống đo lường.',
        'Metadata điều kiện test, quan trọng cho phần Methodology của paper. '
        '"wifi" hoặc "ethernet" cho biết test được thực hiện qua mạng cố định/WiFi (Docker local). '
        'Chỉ hỗ trợ trên Chromium-based browsers (Chrome, Edge, Opera). Safari và Firefox KHÔNG hỗ trợ.'
    )

    doc.add_page_break()

    # ===== PHẦN 3: BẢNG TỔNG HỢP =====
    doc.add_heading('3. Bảng tổng hợp', level=1)

    summary_table = doc.add_table(rows=17, cols=5)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    s_headers = ['#', 'Thông số', 'Công thức / Nguồn', 'Hướng tốt', 'Mức độ quan trọng']
    for i, h in enumerate(s_headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    s_data = [
        ['1', 'Throughput', 'avg(segment speeds 1s)', 'Cao hơn ↑', '★★★ Bắt buộc'],
        ['2', 'TTFB', 'responseStart - requestStart', 'Thấp hơn ↓', '★★★ Bắt buộc'],
        ['3', 'SDT', 'endDate - startDate', 'Thấp hơn ↓', '★★★ Bắt buộc'],
        ['4', 'Jitter', '|SDT_i - SDT_{i-1}|', 'Thấp hơn ↓', '★★☆ Nên có'],
        ['5', 'Buffer', 'getBufferLength()', 'Cao hơn ↑', '★★☆ Nên có'],
        ['6', 'StallCount', 'count(BUFFER_EMPTY)', 'Thấp hơn ↓', '★★★ Bắt buộc'],
        ['7', 'StallDuration', 'sum(EMPTY→LOADED)', 'Thấp hơn ↓', '★★★ Bắt buộc'],
        ['8', 'RebufferingRatio', 'stallMs / playbackMs', 'Thấp hơn ↓', '★★★ Bắt buộc'],
        ['9', 'Bitrate', 'rep.bitrateInKbit', 'Cao hơn ↑', '★★★ Bắt buộc'],
        ['10', 'QualitySwitchCount', 'count(QUALITY_CHANGE)', 'Thấp hơn ↓', '★★☆ Nên có'],
        ['11', 'DroppedFrames', 'droppedVideoFrames', 'Thấp hơn ↓', '★☆☆ Bổ sung'],
        ['12', 'FPS', 'Δframes / Δtime', 'Cao hơn ↑', '★☆☆ Bổ sung'],
        ['13', 'DownloadSpeed', '(bytes×8) / SDT', 'Cao hơn ↑', '★★☆ Nên có'],
        ['14', 'Resolution', 'rep.width × height', '—', 'Metadata'],
        ['15', 'Protocol', 'nextHopProtocol', '—', 'Biến độc lập'],
        ['16', 'NetworkType', 'connection.type', '—', 'Methodology'],
    ]

    for r, row_data in enumerate(s_data):
        for c, val in enumerate(row_data):
            cell = summary_table.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            # Highlight mandatory metrics
            if '★★★' in val:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    legend = doc.add_paragraph()
    legend.add_run('Chú thích: ').bold = True
    legend.add_run('★★★ = Bắt buộc trong paper Q1 | ★★☆ = Nên có | ★☆☆ = Bổ sung thêm | Metadata/Biến độc lập = Không so sánh trực tiếp')

    # ===== PHẦN 4: TÀI LIỆU THAM KHẢO =====
    doc.add_heading('4. Tài liệu tham khảo', level=1)

    refs = [
        '[1] Seufert, M., Egger, S., Slanina, M., Zinner, T., Hoßfeld, T., & Tran-Gia, P. (2015). '
        '"A Survey on Quality of Experience of HTTP Adaptive Streaming." '
        'IEEE Communications Surveys & Tutorials, 17(1), 469-492.',

        '[2] Bentaleb, A., Taani, B., Begen, A. C., Timmerer, C., & Zimmermann, R. (2019). '
        '"A Survey on Bitrate Adaptation Schemes for Streaming Media Over HTTP." '
        'IEEE Communications Surveys & Tutorials, 21(1), 562-585.',

        '[3] Bhat, D., Rizk, A., & Zink, M. (2020). '
        '"Not So QUIC: A Performance Study of DASH over QUIC." '
        'ACM NOSSDAV 2020.',

        '[4] Palmer, M., Krüger, T., Chandaria, B., & Sherrer, S. (2018). '
        '"The QUIC Fix for Optimal Video Streaming." '
        'ACM MMSys 2018.',

        '[5] Yu, J., et al. (2021). '
        '"Can QUIC Replace TCP for Web Video Streaming?" '
        'IEEE INFOCOM Workshop 2021.',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)

    # ===== SAVE =====
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MetricCalculationReport.docx')
    doc.save(output_path)
    print(f'✅ Đã tạo: {output_path}')


if __name__ == '__main__':
    main()
